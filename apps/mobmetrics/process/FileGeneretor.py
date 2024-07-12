from docx import Document
import pandas as pd

from ..metrics.utils.StayPoints import StayPoints

class FileGenerator():
  def __init__(self, trace):
    self.trace = trace

  def export(self):
    self.file = Document()

    self.file.add_heading('MobMetrics', level=1)
    self.file.add_paragraph('Mobility Model metrics analised by individual and globals aspects')

    self.format_trace()

    for id in self.possibleId:
      filtred_trace = self.trace[self.trace['id'] == id]
      individual_trace = filtred_trace.sort_values(by='time')

      self.individualMetric(individual_trace, id) 

    self.file.save('apps/mobmetrics/process/media/metrics.docx')

  def individualMetric(self, individual_trace, id):
    self.file.add_heading(f'Object Id {id}', level=2)

    self.file.add_heading(f'Stay Points', level=3)
    
    self.insert_table(StayPoints().extract(individual_trace, 100, 90))

  def global_metric(self):
    pass

  def format_trace(self):
    self.trace = pd.read_csv(self.trace)

    self.possibleId = sorted(self.trace['id'].unique())

  def insert_table(self, pandas_table):
    table = self.file.add_table(rows=1, cols=len(pandas_table.columns))

    # Adiciona cabeçalhos
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(pandas_table.columns):
        hdr_cells[i].text = column

    # Adiciona as linhas do DataFrame na tabela
    for index, row in pandas_table.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)